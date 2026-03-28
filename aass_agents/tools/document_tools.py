"""
Document ingestion tools for SL Agents.

Supported formats: PDF, DOCX, Markdown, TXT, HTML, XLSX, CSV.
Drop files into the `documents/` folder next to this project,
then pass the filename or full path to these tools.

Tools:
  read_document(path)                  — auto-detect format, return clean text
  read_document_pages(path, pages)     — PDF only: read specific page range
  list_documents(folder)               — list ingestable files in a folder
  search_document(path, query)         — keyword search within a document
"""

from __future__ import annotations

import csv
import io
import os
from pathlib import Path

# Default intake folder — agents can reference files by name without full path
_DOCS_ROOT = Path(__file__).parent.parent / "documents"


def _resolve(path: str) -> Path:
    """Resolve a filename or full path against the documents folder."""
    p = Path(path)
    if p.is_absolute() and p.exists():
        return p
    candidate = _DOCS_ROOT / p
    if candidate.exists():
        return candidate
    # Try as-is (relative to cwd)
    if p.exists():
        return p.resolve()
    raise FileNotFoundError(
        f"File not found: {path}\n"
        f"Place files in the documents/ folder: {_DOCS_ROOT}"
    )


def _read_pdf(path: Path, pages: list[int] | None = None) -> str:
    try:
        import pdfplumber
    except ImportError:
        return "Error: pdfplumber not installed. Run: pip install pdfplumber"

    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        total = len(pdf.pages)
        target = pdf.pages if pages is None else [
            pdf.pages[i - 1] for i in pages if 1 <= i <= total
        ]
        for page in target:
            text = page.extract_text() or ""
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        text += "\n" + " | ".join(str(c or "") for c in row)
            chunks.append(text.strip())
    return f"[PDF: {path.name} — {total} pages]\n\n" + "\n\n---\n\n".join(chunks)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            prefix = ""
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                prefix = "#" * int(level) + " " if level.isdigit() else "## "
            parts.append(prefix + para.text.strip())
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        parts.append("\n".join(rows))
    return f"[DOCX: {path.name}]\n\n" + "\n\n".join(parts)


def _read_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return "Error: beautifulsoup4 not installed. Run: pip install beautifulsoup4"

    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return f"[HTML: {path.name}]\n\n" + "\n".join(lines)


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows.append(" | ".join(str(c or "") for c in row))
        if rows:
            parts.append(f"### Sheet: {sheet}\n" + "\n".join(rows))
    return f"[XLSX: {path.name}]\n\n" + "\n\n".join(parts)


def _read_csv(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(row) for row in reader if any(c.strip() for c in row)]
    return f"[CSV: {path.name}]\n\n" + "\n".join(rows)


def _read_text(path: Path) -> str:
    ext = path.suffix.lower()
    label = {".md": "Markdown", ".txt": "Text", ".rst": "RST"}.get(ext, ext.upper())
    content = path.read_text(encoding="utf-8", errors="replace")
    return f"[{label}: {path.name}]\n\n{content}"


# ── Public tools ──────────────────────────────────────────────────────────────

def read_document(path: str) -> str:
    """
    Read and extract text from a document file.
    Supports: PDF, DOCX, Markdown, TXT, HTML, XLSX, CSV.
    Pass a filename (e.g. 'report.pdf') or full path.
    Files in the documents/ folder can be referenced by name alone.
    Returns extracted text with format label.
    """
    try:
        p = _resolve(path)
    except FileNotFoundError as e:
        return f"Error: {e}"

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _read_pdf(p)
        elif ext in (".docx", ".doc"):
            return _read_docx(p)
        elif ext in (".html", ".htm"):
            return _read_html(p)
        elif ext in (".xlsx", ".xls"):
            return _read_xlsx(p)
        elif ext == ".csv":
            return _read_csv(p)
        elif ext in (".md", ".txt", ".rst", ".text"):
            return _read_text(p)
        else:
            # Try as plain text fallback
            return _read_text(p)
    except Exception as exc:
        return f"Error reading {p.name}: {exc}"


def read_document_pages(path: str, pages: str) -> str:
    """
    Read specific pages from a PDF document.
    path: filename or full path to the PDF.
    pages: comma-separated page numbers or ranges, e.g. '1,3,5-8'.
    Returns extracted text from the specified pages only.
    """
    try:
        p = _resolve(path)
    except FileNotFoundError as e:
        return f"Error: {e}"

    if p.suffix.lower() != ".pdf":
        return f"Error: read_document_pages only supports PDF files. Got: {p.suffix}"

    # Parse page spec: "1,3,5-8" → [1, 3, 5, 6, 7, 8]
    page_nums: list[int] = []
    for part in pages.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            try:
                page_nums.extend(range(int(start.strip()), int(end.strip()) + 1))
            except ValueError:
                return f"Error: invalid page range '{part}'"
        else:
            try:
                page_nums.append(int(part))
            except ValueError:
                return f"Error: invalid page number '{part}'"

    try:
        return _read_pdf(p, pages=page_nums)
    except Exception as exc:
        return f"Error reading {p.name}: {exc}"


def list_documents(folder: str = "") -> str:
    """
    List all ingestable documents in a folder.
    folder: path to scan. Defaults to the documents/ intake folder.
    Returns a formatted list of files with sizes.
    """
    SUPPORTED = {".pdf", ".docx", ".doc", ".html", ".htm",
                 ".xlsx", ".xls", ".csv", ".md", ".txt", ".rst"}

    target = Path(folder) if folder else _DOCS_ROOT
    if not target.exists():
        return f"Folder not found: {target}"

    files = sorted(
        [f for f in target.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED],
        key=lambda f: f.stat().st_mtime, reverse=True
    )

    if not files:
        return f"No documents found in: {target}\nSupported: {', '.join(sorted(SUPPORTED))}"

    lines = [f"Documents in {target} ({len(files)} files):\n"]
    for f in files:
        size = f.stat().st_size
        size_str = f"{size:,} bytes" if size < 1024 else f"{size // 1024:,} KB"
        rel = f.relative_to(target) if f.is_relative_to(target) else f
        lines.append(f"  {rel}  [{f.suffix.upper()[1:]}]  {size_str}")
    return "\n".join(lines)


def search_document(path: str, query: str) -> str:
    """
    Search for a keyword or phrase within a document.
    path: filename or full path.
    query: text to search for (case-insensitive).
    Returns matching lines with context (±2 lines).
    """
    content = read_document(path)
    if content.startswith("Error:"):
        return content

    lines = content.splitlines()
    q = query.lower()
    matches = []
    for i, line in enumerate(lines):
        if q in line.lower():
            start = max(0, i - 2)
            end = min(len(lines), i + 3)
            snippet = "\n".join(
                f"{'>>>' if j == i else '   '} {lines[j]}"
                for j in range(start, end)
            )
            matches.append(f"[Line {i + 1}]\n{snippet}")

    if not matches:
        return f"No matches for '{query}' in {Path(path).name}"
    return f"Found {len(matches)} match(es) for '{query}' in {Path(path).name}:\n\n" + \
           "\n\n".join(matches[:20])  # cap at 20 matches
